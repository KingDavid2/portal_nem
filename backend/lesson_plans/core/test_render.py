"""The docx must mirror the Planeabot example layout: Datos table, per-momento
Duración paragraphs, Paso/Dinámica tables, rubric table.

Ported from the M1 spike's `test_render_docx.py`, adapted for the BytesIO
return value (design D1a).
"""

from __future__ import annotations

from docx import Document

from lesson_plans.core.render.docx import render_docx
from lesson_plans.core.render.markdown import render_md
from lesson_plans.core.schema import (
    ArticulatingAxis,
    ContentPda,
    Datos,
    Momento,
    Proyecto,
    Rubric,
    RubricCriterion,
    Session,
    Stage,
    Step,
)


def _proyecto() -> Proyecto:
    return Proyecto(
        datos=Datos(
            school_name="Escuela Secundaria General Jesús Reyes Heroles",
            grade="SEGUNDO",
            campo_formativo="ÉTICA, NATURALEZA Y SOCIEDADES",
            date="JUEVES, 18 DE SEPTIEMBRE DE 2025",
        ),
        title="Héroes y Gestas",
        purpose="Fomentar la comprensión de los movimientos de resistencia.",
        articulating_axes=[ArticulatingAxis(name="Inclusión", justification="Integrar a todos.")],
        problem_or_theme="LA INDEPENDENCIA DE MEXICO",
        contents_and_pdas=[
            ContentPda(content="Las gestas de resistencia.", pdas=["Relaciona la Revolución de 1810."])
        ],
        stages=[
            Stage(
                name="Planeación",
                momentos=[
                    Momento(
                        number=1,
                        name="Identificación",
                        sessions=[
                            Session(
                                duration_minutes=30,
                                steps=[
                                    Step(number=1, dynamic="Introducción a la Independencia."),
                                    Step(number=2, dynamic="Planteamiento del problema."),
                                ],
                            )
                        ],
                    )
                ],
            )
        ],
        rubric=Rubric(
            criteria=[RubricCriterion(criterion="Participación", levels=["Excelente", "Bueno", "Regular", "Inicial"])]
        ),
    )


def _rows(table):
    return [[c.text for c in row.cells] for row in table.rows]


def test_datos_table_is_first_and_key_value_shaped():
    buffer = render_docx(_proyecto())
    doc = Document(buffer)
    datos = _rows(doc.tables[0])
    assert datos[0] == ["CAMPOS DEL PROYECTO", "DATOS"]
    flat = {r[0]: r[1] for r in datos[1:]}
    assert flat["NOMBRE DE LA ESCUELA"] == "Escuela Secundaria General Jesús Reyes Heroles"
    assert flat["FASE"] == "6"
    assert flat["GRADO"] == "SEGUNDO"
    assert flat["METODOLOGÍA DEL PROYECTO"] == "APRENDIZAJE BASADO EN PROYECTOS COMUNITARIOS"
    assert flat["CLAVE CENTRO DE TRABAJO"] == "[ESPACIO PARA SER LLENADO POR EL DOCENTE]"


def _datos_map(proyecto):
    datos = _rows(Document(render_docx(proyecto)).tables[0])
    return {r[0]: r[1] for r in datos[1:]}


def test_datos_table_carries_the_subject_and_the_period():
    proyecto = _proyecto()
    proyecto.datos = proyecto.datos.model_copy(
        update={"subject": "Español", "start_date": "2026-01-12", "end_date": "2026-02-09"}
    )

    flat = _datos_map(proyecto)

    assert flat["ASIGNATURA"] == "Español"
    assert flat["PERIODO"] == "2026-01-12 a 2026-02-09"


def test_datos_table_omits_the_subject_and_period_when_unknown():
    """Plans generated before the header carried them must render as before."""
    flat = _datos_map(_proyecto())

    assert "ASIGNATURA" not in flat
    assert "PERIODO" not in flat


def test_render_md_carries_the_subject_and_the_period():
    proyecto = _proyecto()
    proyecto.datos = proyecto.datos.model_copy(
        update={"subject": "Español", "start_date": "2026-01-12", "end_date": "2026-02-09"}
    )

    md = render_md(proyecto)

    assert "**Asignatura:** Español" in md
    assert "**Periodo:** 2026-01-12 a 2026-02-09" in md


def test_render_md_omits_the_subject_and_period_when_unknown():
    md = render_md(_proyecto())

    assert "Asignatura" not in md
    assert "Periodo" not in md


def test_headings_and_duracion_present():
    buffer = render_docx(_proyecto())
    doc = Document(buffer)
    texts = [p.text for p in doc.paragraphs]
    assert "Contenidos y PDAs seleccionados" in texts
    assert "Fase 1: Planeación" in texts
    assert "Momento 1: Identificación" in texts
    assert "Duración: 30 minutos" in texts
    assert "Relaciona la Revolución de 1810." in texts


def test_paso_dinamica_table_present():
    buffer = render_docx(_proyecto())
    doc = Document(buffer)
    all_tables = [_rows(t) for t in doc.tables]
    paso = next(t for t in all_tables if t and t[0] == ["Paso", "Dinámica"])
    assert paso[1] == ["1", "Introducción a la Independencia."]
    assert paso[2] == ["2", "Planteamiento del problema."]


def test_rubric_table_has_criterion_and_four_levels():
    buffer = render_docx(_proyecto())
    doc = Document(buffer)
    all_tables = [_rows(t) for t in doc.tables]
    rubric = next(t for t in all_tables if t and t[0][0] == "Criterio")
    assert rubric[1] == ["Participación", "Excelente", "Bueno", "Regular", "Inicial"]


def test_render_md_includes_title_and_pdas():
    md = render_md(_proyecto())
    assert md.startswith("# Héroes y Gestas")
    assert "Relaciona la Revolución de 1810." in md
    assert "Duración:** 30 minutos" in md
