"""Render a Proyecto to a .docx that mirrors the Planeabot example layout.

Structure (verified against designs/examples/Ejemplo proyecto ...docx):
  1. Datos — one 2-col key/value table (header row "CAMPOS DEL PROYECTO" | "DATOS").
  2. "Contenidos y PDAs seleccionados" heading + one paragraph per contenido/PDA.
  3. Per stage: "Fase N: <nombre>" -> per momento "Momento N: <nombre>" -> per
     session a "Duración: N minutos" paragraph + a 2-col Paso/Dinámica table.
  4. Rubric — criteria x 4 achievement levels table.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from ..schema import Proyecto

_DISCLAIMER = (
    "IMPORTANTE: Este material ha sido generado con inteligencia artificial como apoyo "
    "educativo y debe ser revisado y adaptado según las necesidades específicas del docente "
    "y del contexto educativo."
)


def _datos_rows(p: Proyecto) -> list[tuple[str, str]]:
    ejes = "\n".join(f"- {a.name}: {a.justification}" for a in p.articulating_axes)
    d = p.datos
    return [
        ("NOMBRE DE LA ESCUELA", d.school_name),
        ("CLAVE CENTRO DE TRABAJO", d.cct),
        ("FASE", str(d.phase)),
        ("GRADO", d.grade),
        ("TÍTULO DEL PROYECTO", p.title),
        ("PROPÓSITO DEL PROYECTO", p.purpose),
        ("VINCULACIÓN CON LOS EJES ARTICULADORES", ejes),
        ("CAMPO FORMATIVO", d.campo_formativo),
        ("METODOLOGÍA DEL PROYECTO", d.methodology),
        ("PROBLEMÁTICA O TEMA", p.problem_or_theme),
        ("FECHA", d.date),
    ]


def render_docx(proyecto: Proyecto, path: str | Path) -> Path:
    doc = Document()
    doc.add_paragraph(_DISCLAIMER)

    # 1. Datos table.
    rows = _datos_rows(proyecto)
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "CAMPOS DEL PROYECTO"
    table.rows[0].cells[1].text = "DATOS"
    for i, (key, value) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value

    # 2. Contenidos y PDAs.
    doc.add_heading("Contenidos y PDAs seleccionados", level=1)
    for group in proyecto.contents_and_pdas:
        doc.add_paragraph(group.content)
        for pda in group.pdas:
            doc.add_paragraph(pda)

    # 3. Stages -> momentos -> sessions.
    for stage_idx, stage in enumerate(proyecto.stages, start=1):
        doc.add_heading(f"Fase {stage_idx}: {stage.name}", level=1)
        for momento in stage.momentos:
            doc.add_heading(f"Momento {momento.number}: {momento.name}", level=2)
            for session in momento.sessions:
                doc.add_paragraph(f"Duración: {session.duration_minutes} minutos")
                st = doc.add_table(rows=len(session.steps) + 1, cols=2)
                st.style = "Table Grid"
                st.rows[0].cells[0].text = "Paso"
                st.rows[0].cells[1].text = "Dinámica"
                for j, step in enumerate(session.steps, start=1):
                    st.rows[j].cells[0].text = str(step.number)
                    st.rows[j].cells[1].text = step.dynamic

    # 4. Rubric.
    doc.add_heading("Rúbrica de evaluación", level=1)
    rt = doc.add_table(rows=len(proyecto.rubric.criteria) + 1, cols=5)
    rt.style = "Table Grid"
    header = ["Criterio", "Nivel 1", "Nivel 2", "Nivel 3", "Nivel 4"]
    for col, label in enumerate(header):
        rt.rows[0].cells[col].text = label
    for i, crit in enumerate(proyecto.rubric.criteria, start=1):
        rt.rows[i].cells[0].text = crit.criterion
        for col, level in enumerate(crit.levels[:4], start=1):
            rt.rows[i].cells[col].text = level

    path = Path(path)
    doc.save(str(path))
    return path
