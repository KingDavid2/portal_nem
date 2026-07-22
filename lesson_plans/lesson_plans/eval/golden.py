"""The golden reference: the Planeabot example proyecto, as plain text.

Extracted from the repo's example .docx at read time (python-docx is already a
dependency) so the judge scores candidates against a real teacher-accepted artifact.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

# lesson_plans/lesson_plans/eval/golden.py -> repo root is four parents up.
_REPO_ROOT = Path(__file__).resolve().parents[3]
GOLDEN_DOCX = _REPO_ROOT / "designs" / "examples" / "Ejemplo proyecto ETICA NATURALEZA Y SOCIEDADES.docx"


def golden_text() -> str:
    """Flatten the example docx (paragraphs + table cells) to reference text."""
    doc = Document(str(GOLDEN_DOCX))
    lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)
